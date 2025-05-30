from typing import Annotated, Any, Dict
from semantic_kernel.functions import kernel_function
from .base_plugin import BasePlugin

from docx import Document

class WordPlugin(BasePlugin):
    
    def __init__(self, settings : Dict= {}):

        super().__init__(settings)

        self.location = settings.get("location","")
        self.location_type = settings.get("location_type","blob")
        self.location_path = settings.get("location_path","/")

    def download_file(self, file_name: str) -> str:
        """
        Download files from the specified location.
        """
        # Implement the logic to download files from the specified location
        
        # For example, if using Azure Blob Storage:
        blob_client = BlobClient.from_blob_url(f"https://{self.location}.blob.core.windows.net/{self.location_path}/{file_name}")
        with open(file_name, "wb") as file:
           blob_data = blob_client.download_blob()
           file.write(blob_data.readall())

    def create_from_template(self, 
                     template_file_path: str, 
                     output_file_path: str, 
                     variables: Dict[str, str]) -> None:
        """
        Create a Word document from a template with specified replacements.
        """
        # Implement the logic to create a Word document from a template
        
        variables = {
            "${EMPLOEE_NAME}": "Example Name",
            "${EMPLOEE_TITLE}": "Software Engineer",
            "${EMPLOEE_ID}": "302929393",
            "${EMPLOEE_ADDRESS}": "דרך השלום מנחם בגין דוגמא",
            "${EMPLOEE_PHONE}": "+972-5056000000",
            "${EMPLOEE_EMAIL}": "example@example.com",
            "${START_DATE}": "03 Jan, 2021",
            "${SALARY}": "10,000",
            "${SALARY_30}": "3,000",
            "${SALARY_70}": "7,000",
        }

        template_document = Document(template_file_path)

        for variable_key, variable_value in variables.items():
            for paragraph in template_document.paragraphs:
                self.replace_text_in_paragraph(paragraph, variable_key, variable_value)

            for table in template_document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            self.replace_text_in_paragraph(paragraph, variable_key, variable_value)

        template_document.save(output_file_path)

    def replace_text_in_paragraph(self,paragraph, key, value):
        if key in paragraph.text:
            inline = paragraph.runs
            for item in inline:
                if key in item.text:
                    item.text = item.text.replace(key, value)
